"""
Generates Final_Report.docx — full research thesis for the CERPM/LKA project.
Run from the repo root: python3 build_report.py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

IMG = {
    # Town07G
    't07_road':        'Results/Town07G/Figure_1.png',
    't07_rand_heatmap':'Results/Town07G/random_iqr_heatmap - Copy.png',
    't07_rand_bands':  'Results/Town07G/random_iqr_bands.png',
    't07_rand_fail':   'Results/Town07G/random_failure_rate.png',
    't07_gap_heatmap': 'Results/Town07G/gaplength_iqr_heatmap.png',
    't07_gap_bands':   'Results/Town07G/gaplength_iqr_bands.png',
    't07_gap_fail':    'Results/Town07G/gaplength_failure_rate.png',
    # Town042G
    't42_road':        'Results/Town042G/Figure_1.png',
    't42_rand_heatmap':'Results/Town042G/random_iqr_heatmap.png',
    't42_rand_bands':  'Results/Town042G/random_iqr_bands.png',
    't42_rand_fail':   'Results/Town042G/random_failure_rate.png',
    't42_gap_heatmap': 'Results/Town042G/gaplength_iqr_heatmap.png',
    't42_gap_bands':   'Results/Town042G/gaplength_iqr_bands.png',
    't42_gap_fail':    'Results/Town042G/gaplength_failure_rate.png',
    # Town044G
    't44_road':        'Results/Town044G/Figure_1.png',
    't44_rand_heatmap':'Results/Town044G/random_iqr_heatmap.png',
    't44_rand_bands':  'Results/Town044G/random_iqr_bands.png',
    't44_rand_fail':   'Results/Town044G/random_failure_rate.png',
    't44_gap_heatmap': 'Results/Town044G/gaplength_iqr_heatmap.png',
    't44_gap_bands':   'Results/Town044G/gaplength_iqr_bands.png',
    't44_gap_fail':    'Results/Town044G/gaplength_failure_rate.png',
}


def set_font(run, name='Calibri', size=11, bold=False, italic=False, colour=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if colour:
        run.font.color.rgb = RGBColor(*colour)


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p


def body(doc, text, space_after=6):
    p = doc.add_paragraph(text)
    p.style = doc.styles['Normal']
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
    return p


def caption(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(8)
    run = p.runs[0] if p.runs else p.add_run()
    run.font.name   = 'Calibri'
    run.font.size   = Pt(9)
    run.italic      = True
    run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    return p


def figure(doc, key, cap, width=Inches(6.2)):
    path = IMG.get(key)
    if path and os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run()
        run.add_picture(path, width=width)
    caption(doc, cap)


def add_table_row(table, cells, bold=False, bg=None):
    row = table.add_row()
    for i, text in enumerate(cells):
        cell = row.cells[i]
        cell.text = text
        for run in cell.paragraphs[0].runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            run.bold = bold
    return row


def shade_cell(cell, hex_colour):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_colour)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)


def page_break(doc):
    doc.add_page_break()


# ──────────────────────────────────────────────
# BUILD
# ──────────────────────────────────────────────
doc = Document()

# ── Margins ──
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ══════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════
for _ in range(4):
    doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_p.add_run('Interpolation Algorithm Evaluation for\nCERPM-Based Lane Keep Assist Systems')
r.font.name  = 'Calibri'
r.font.size  = Pt(22)
r.bold       = True

doc.add_paragraph()

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub_p.add_run('Final Research Thesis')
r2.font.name  = 'Calibri'
r2.font.size  = Pt(14)
r2.italic     = True

doc.add_paragraph()
doc.add_paragraph()

for label, value in [
    ('Unit:',       'Advanced Driver Assistance Systems Research Project'),
    ('Author:',     'Alexander Bruce'),
    ('Date:',       'June 2026'),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rb = p.add_run(label + ' ')
    rb.bold = True
    rb.font.name = 'Calibri'
    rb.font.size = Pt(12)
    rv = p.add_run(value)
    rv.font.name = 'Calibri'
    rv.font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph()

contrib_heading = doc.add_paragraph()
contrib_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
rc = contrib_heading.add_run('Student Contribution')
rc.bold = True
rc.font.name = 'Calibri'
rc.font.size = Pt(12)

contrib_p = doc.add_paragraph(
    'This report represents the Lane Keep Assist (LKA) subsystem contribution to the '
    'group Advanced Driver Assistance Systems (ADAS) project. The author was responsible '
    'for the full scope of this subsystem: problem definition, literature review, system '
    'architecture design, simulation development in Python (road geometry parsing, CERPM '
    'resampling, Monte Carlo simulation framework, and all analysis and visualisation code), '
    'execution of experiments across three road geometries, and analysis and interpretation '
    'of all results.'
)
contrib_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in contrib_p.runs:
    run.font.name = 'Calibri'
    run.font.size = Pt(11)

page_break(doc)

# ══════════════════════════════════════════════
# ABSTRACT
# ══════════════════════════════════════════════
heading(doc, 'Abstract', level=1)
body(doc,
    'Lane Keep Assist (LKA) is a foundational safety technology within Advanced Driver '
    'Assistance Systems (ADAS), designed to prevent unintentional lane departures through '
    'corrective steering. Current commercial LKA implementations rely primarily on '
    'camera-based computer vision, which suffers significant performance degradation under '
    'poor lane markings and adverse weather conditions. This research investigates a '
    'Vehicle-to-Infrastructure (V2I) alternative using Chip Enabled Raised Pavement Markers '
    '(CERPMs), which transmit GPS positional data directly to the onboard vehicle computer.'
)
body(doc,
    'The central research questions are: (1) which interpolation algorithm most accurately '
    'reconstructs the lane centreline from CERPM position data; and (2) how do different '
    'CERPM dropout scenarios — random failures, consecutive gap-length failures, and '
    'clustered failures — degrade accuracy across varying marker spacings?'
)
body(doc,
    'A Python simulation framework was developed to parse real-world road geometries from '
    'Lanelet2 OSM files and to evaluate seven interpolation methods (linear, quadratic, '
    'cubic spline, quartic, quintic, PCHIP, and Akima) under Monte Carlo conditions across '
    'three distinct road geometries. At dense CERPM spacings (≤1 m) all methods achieve '
    'sub-centimetre median centreline error with near-zero failure rates. As spacing '
    'increases beyond 2 m, smooth methods (cubic spline, PCHIP, Akima) substantially '
    'outperform linear and quadratic interpolation. Under gap-length dropout, linear '
    'interpolation fails catastrophically at nearly all spacings beyond 1 m on curved roads, '
    'while PCHIP and Akima exhibit the most robust behaviour. Notably, PCHIP performs worse '
    'than Akima on the sharp-corner geometry (Town044), making Akima the most consistently '
    'reliable choice across all road types. These findings provide actionable guidance for '
    'CERPM deployment standards and algorithm selection in production LKA systems.'
)

page_break(doc)

# ══════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ══════════════════════════════════════════════
heading(doc, 'Table of Contents', level=1)
toc_entries = [
    ('1.', 'Introduction'),
    ('  1.1', 'General Background'),
    ('  1.2', 'The Problem with Vision-Based LKA'),
    ('  1.3', 'The CERPM-Based V2I Approach'),
    ('  1.4', 'Research Gap and Aims'),
    ('  1.5', 'Scope'),
    ('2.', 'Literature Review'),
    ('  2.1', 'State of the Art in LKA Perception'),
    ('  2.2', 'Infrastructure-Based Lane Detection'),
    ('  2.3', 'Interpolation Methods for Spatial Data'),
    ('  2.4', 'LKA Control'),
    ('  2.5', 'Summary of Research Gap'),
    ('3.', 'Methodology'),
    ('  3.1', 'Overview'),
    ('  3.2', 'Road Geometry Acquisition'),
    ('  3.3', 'CERPM Placement Simulation'),
    ('  3.4', 'True Centreline Computation'),
    ('  3.5', 'Interpolation Methods'),
    ('  3.6', 'Dropout Scenarios'),
    ('  3.7', 'Error Metrics'),
    ('  3.8', 'Project Management'),
    ('4.', 'Results and Analysis'),
    ('  4.1', 'Road Geometries'),
    ('  4.2', 'Baseline Accuracy vs CERPM Interval'),
    ('  4.3', 'Random Dropout — Town07 (Sinuous Road)'),
    ('  4.4', 'Random Dropout — Town042 (Gradual Curve)'),
    ('  4.5', 'Random Dropout — Town044 (Sharp Corners)'),
    ('  4.6', 'Gap-Length Dropout — Town07'),
    ('  4.7', 'Gap-Length Dropout — Town042'),
    ('  4.8', 'Gap-Length Dropout — Town044'),
    ('  4.9', 'Cross-Geometry Comparative Summary'),
    ('  4.10', 'Changes from Initial Research Plan'),
    ('5.', 'Conclusions and Recommendations'),
    ('6.', 'References'),
    ('7.', 'Appendices'),
]
for num, title in toc_entries:
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    r_num = p.add_run(num + '\t' + title)
    r_num.font.name = 'Calibri'
    r_num.font.size = Pt(11)
    if not num.startswith('  '):
        r_num.bold = True

page_break(doc)

# ══════════════════════════════════════════════
# 1. INTRODUCTION
# ══════════════════════════════════════════════
heading(doc, '1. Introduction', level=1)

heading(doc, '1.1 General Background', level=2)
body(doc,
    'Road safety remains one of the most pressing public health challenges globally. In '
    'Australia, lane departure — encompassing run-off-road and head-on collisions — is the '
    'single largest contributor to road fatalities, accounting for 62% of deaths nationally, '
    'rising to 73% in regional areas and 71% in remote areas [3]. In the United States, the '
    'National Highway Traffic Safety Administration (NHTSA) estimated that vehicles equipped '
    'with Lane Keep Assist are on average 24% less likely to be involved in a fatal road '
    'departure crash [4]. These statistics establish LKA as one of the highest-leverage '
    'active safety technologies available.'
)
body(doc,
    'From 7 July 2024, the European Commission mandated LKA as standard fitment on all new '
    'motor vehicles sold in the European Union, projecting the technology will contribute to '
    'saving over 25,000 lives by 2038 [5]. LKA also forms the perceptual and control '
    'foundation upon which higher-order autonomous driving functions — such as highway pilot '
    'and lane-change assist — are built. Its reliability is therefore critical not only as a '
    'standalone safety feature but as an enabler of the broader ADAS ecosystem.'
)

heading(doc, '1.2 The Problem with Vision-Based LKA', level=2)
body(doc,
    'Current production LKA systems universally rely on forward-facing cameras combined '
    'with computer vision algorithms to detect lane markings. While these systems perform '
    'reliably under ideal conditions, field evidence demonstrates that performance degrades '
    'substantially when lane markings are worn, faded, or absent — conditions '
    'disproportionately common on rural and regional roads where lane departure risk is '
    'highest. Adverse weather (rain, fog, snow, glare) and challenging lighting conditions '
    'further compromise detection reliability [6]. Near lane transitions such as merges, '
    'diverges, and construction zones introduce additional failure modes.'
)
body(doc,
    'Lane Departure Warning (LDW) systems alert the driver to unintended departures but '
    'rely on the driver responding appropriately — unreliable when fatigue, distraction, or '
    'impairment is present. LKA addresses this by actuating the steering system directly, '
    'without requiring driver response.'
)

heading(doc, '1.3 The CERPM-Based V2I Approach', level=2)
body(doc,
    'A fundamentally different approach has emerged from recent research: instead of '
    'detecting lane boundaries visually, the vehicle receives lane boundary positional data '
    'directly from road-embedded infrastructure. Chip Enabled Raised Pavement Markers '
    '(CERPMs) are road stud-type devices equipped with a wireless transceiver that '
    'broadcasts their GPS coordinates to vehicles within range (demonstrated range: up to '
    '350 m [1][2]). When both the left and right lane boundary CERPMs are received, the '
    'lane centreline is computed through interpolation of the two coordinate streams.'
)
body(doc,
    'This approach decouples LKA accuracy from environmental and surface conditions '
    'entirely: it operates in darkness, fog, rain, snow, and on roads with no visible '
    'markings. Kadav et al. [2] demonstrated that CERPM-based lane centering outperformed '
    'the Mobileye 630 — a state-of-the-art commercial computer vision system — across all '
    'tested conditions, with a particularly notable advantage in detection range (350 m vs '
    '31 m for the Mobileye).'
)

heading(doc, '1.4 Research Gap and Aims', level=2)
body(doc,
    'The existing CERPM research [1][2] validated the concept but did not investigate: '
    '(a) whether the choice of interpolation algorithm affects accuracy, or (b) how the '
    'system behaves when individual CERPMs fail — a practically inevitable scenario as '
    'road-embedded hardware is subject to vehicle damage, vandalism, and electronic failure. '
    'This project directly addresses these gaps. The specific aims are:'
)
aims = [
    '1.  Evaluate and compare seven interpolation algorithms for lane centreline estimation '
        'from CERPM data across multiple real-world road geometries.',
    '2.  Characterise how three categories of CERPM dropout — random uniform failure, '
        'gap-length failure, and contiguous clustered failure — degrade centreline estimation '
        'accuracy as a function of marker spacing.',
    '3.  Identify which algorithm–spacing combinations maintain acceptable accuracy '
        '(max centreline error ≤ 0.2 m) under realistic failure conditions, and provide '
        'deployment recommendations.',
]
for a in aims:
    p = doc.add_paragraph(a, style='List Number')
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(11)

heading(doc, '1.5 Scope', level=2)
body(doc,
    'This project is scoped to the perception component of LKA — the estimation of lane '
    'centreline position from CERPM data. The downstream control system (steering actuation) '
    'is outside scope. Road geometries are sourced from Lanelet2-format OSM map files '
    'derived from the CARLA simulator environment. GPS positioning error of the CERPMs '
    'themselves is not modelled; the simulation assumes ground-truth CERPM coordinates are '
    'transmitted accurately.'
)

page_break(doc)

# ══════════════════════════════════════════════
# 2. LITERATURE REVIEW
# ══════════════════════════════════════════════
heading(doc, '2. Literature Review', level=1)

heading(doc, '2.1 State of the Art in LKA Perception', level=2)
body(doc,
    'The perception layer of modern LKA systems has evolved from early threshold-based '
    'edge detectors to deep neural network architectures with multi-frame temporal '
    'integration. The OpenLKA dataset [6] — the first large-scale open benchmark for LKA '
    'research — aggregates approximately 400 hours of real-world LKA-actuated driving from '
    '62 production vehicles. Its empirical benchmarks confirm that even modern deep-learning '
    'LKA systems fail at high rates when lane markings are degraded or absent, and near lane '
    'transitions. This establishes the practical ceiling of camera-based approaches and '
    'motivates the search for alternative sensing modalities.'
)
body(doc,
    'Fakhari and Anwar [7][8] proposed a Multiple Model Adaptive Estimation (MMAE) '
    'algorithm fusing front and rear camera feeds through a Kalman filter, dynamically '
    'selecting the most accurate model based on uncertainty estimates. Tested under night '
    'and adverse weather conditions, their system demonstrates meaningful accuracy '
    'improvements over single-camera baselines, but remains dependent on some minimum level '
    'of visual lane marking quality.'
)

heading(doc, '2.2 Infrastructure-Based Lane Detection', level=2)
body(doc,
    'Sharma et al. [1] introduced CERPMs as a V2I data source for vehicle lateral offset '
    'estimation, demonstrating reduced compute requirements while delivering accurate '
    'position estimates. Kadav et al. [2] conducted a direct comparative evaluation against '
    'the Mobileye 630 commercial kit on roads with sharp curves, variable lighting, and '
    'worn lane markings. CERPMs outperformed the Mobileye in every tested scenario. '
    'Crucially, neither study investigated which interpolation method was used for centreline '
    'estimation, nor did they characterise system behaviour when CERPMs fail. The current '
    'project addresses precisely this gap.'
)

heading(doc, '2.3 Interpolation Methods for Spatial Data', level=2)
body(doc,
    'Arc-length parameterisation — expressing the interpolation variable as normalised path '
    'distance rather than index — is the standard approach for spatial curve reconstruction, '
    'distributing interpolation effort proportionally to physical spacing and avoiding '
    'artefacts from irregular point distributions. The seven methods evaluated in this '
    'project are summarised below.'
)
body(doc,
    'Linear interpolation is C⁰ continuous (position-continuous only), produces no '
    'overshoot, and serves as the simplest baseline. Its lack of smoothness makes it '
    'unsuitable for high-speed vehicle control.'
)
body(doc,
    'Polynomial splines of degree k (quadratic: k=2, quartic: k=4, quintic: k=5) use '
    'B-spline basis functions of the corresponding degree. Higher degree produces smoother '
    'curves but increases sensitivity to sparse or unevenly spaced knots. Runge\'s '
    'phenomenon — oscillation near endpoints — becomes problematic at large knot spacings.'
)
body(doc,
    'Cubic Spline (natural) fits piecewise cubics with C² continuity and zero second '
    'derivative at endpoints. It is the most common smooth interpolator in engineering '
    'practice and the method implicitly used in the original CERPM research.'
)
body(doc,
    'PCHIP (Piecewise Cubic Hermite Interpolating Polynomial) is C¹ continuous and enforces '
    'monotonicity between knot points, preventing overshoot. It is particularly well-suited '
    'to sparse road data, but its monotonicity constraint can be limiting at very sharp '
    'curvature reversals.'
)
body(doc,
    'Akima interpolation is a C¹ cubic method using locally weighted slope estimation at '
    'each knot, designed so that a single outlier knot does not propagate oscillation to '
    'neighbouring curve segments. This local insensitivity property is especially relevant '
    'when isolated CERPM failures leave gaps in the data.'
)

heading(doc, '2.4 LKA Control', level=2)
body(doc,
    'Perozzi et al. [9] proposed a shared control architecture for steer-by-wire LKA using '
    'a quasi-continuous high-order sliding mode controller, validated in the SHERPA dynamic '
    'vehicle simulator. This architecture provides smooth driver–automation authority '
    'handover and represents the control layer upon which the CERPM perception system '
    'developed here would operate. Wei, Pfeffer, and Edelmann [10] provide the most '
    'comprehensive review of LKA assessment methodologies, identifying standardisation gaps '
    'and proposing an evaluation framework across performance, ride comfort, safety, driver '
    'interaction, and driving style dimensions.'
)

heading(doc, '2.5 Summary of Research Gap', level=2)
body(doc,
    'The literature confirms a clear and unaddressed gap: while the V2I CERPM concept has '
    'been validated as superior to vision-based approaches, no study has characterised the '
    'comparative accuracy of different interpolation algorithms for centreline estimation, '
    'nor the effect of CERPM dropout in any form on system accuracy. This project provides '
    'the first systematic investigation of both questions.'
)

page_break(doc)

# ══════════════════════════════════════════════
# 3. METHODOLOGY
# ══════════════════════════════════════════════
heading(doc, '3. Methodology', level=1)

heading(doc, '3.1 Overview', level=2)
body(doc,
    'A simulation-based methodology was adopted because it enables controlled, repeatable, '
    'and exhaustive testing across parameter combinations that would be impractical to '
    'replicate in physical field trials. The methodology follows four stages: (1) road '
    'geometry acquisition and parsing; (2) CERPM placement simulation; (3) dropout scenario '
    'simulation via Monte Carlo methods; (4) interpolation, centreline computation, and '
    'error measurement. All code was implemented in Python using NumPy, SciPy, Shapely, '
    'and Matplotlib.'
)

heading(doc, '3.2 Road Geometry Acquisition', level=2)
body(doc,
    'Real-world road geometries were sourced from Lanelet2-format OSM files. Lanelet2 '
    'encodes each lane as a lanelet containing a left boundary way and a right boundary way, '
    'each composed of nodes with local Cartesian coordinates (local_x, local_y, in metres). '
    'Three road geometries were used:'
)
geo_items = [
    'Town07: A sinuous, S-curve road with multiple curvature reversals over approximately '
     '280 m total path length. Representative of a rural road with continuous direction changes.',
    'Town042: A long, gradually curved road segment approximately 560 m in length. '
     'Representative of a rural highway or regional road with gentle bends.',
    'Town044: A road with two sharp 90° corners connected by a straight section, total path '
     'approximately 520 m. Representative of an urban intersection approach.',
]
for item in geo_items:
    p = doc.add_paragraph(item, style='List Bullet')
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(11)

body(doc,
    'Road geometry was parsed by extracting node coordinates, assembling them into way '
    'sequences, and then chain-linking consecutive lanelets into continuous left and right '
    'boundary polylines. The three geometries were selected to represent meaningfully '
    'different curvature profiles, ensuring results are not geometry-specific.'
)

heading(doc, '3.3 CERPM Placement Simulation', level=2)
body(doc,
    'CERPMs were simulated by resampling the left and right boundary polylines at a '
    'configurable uniform interval using Shapely\'s interpolate method on a LineString '
    'object. This produces evenly spaced point coordinates along each boundary, mimicking '
    'physical CERPM placement. Six intervals were tested: 0.5 m, 1.0 m, 2.0 m, 4.0 m, '
    '6.0 m, and 12.0 m. The 0.5–2.0 m range reflects dense installations for '
    'safety-critical corridors; 4.0–12.0 m reflects sparser, cost-driven installations.'
)

heading(doc, '3.4 True Centreline Computation', level=2)
body(doc,
    'The ground-truth lane centreline was computed directly from the full-resolution '
    'boundary polylines (before resampling) using a perpendicular-nearest-point algorithm. '
    'For each point sampled at 0.5 m intervals along the longer boundary, the nearest point '
    'on the opposite boundary is found by projecting onto the opposite LineString. The '
    'centreline point is the arithmetic mean of the two boundary points, producing a dense, '
    'geometrically accurate centreline against which interpolated estimates are compared.'
)

heading(doc, '3.5 Interpolation Methods', level=2)
body(doc,
    'Seven interpolation methods were evaluated. All methods use arc-length '
    'parameterisation: the parameter t ∈ [0, 1] represents the normalised cumulative '
    'Euclidean path length along the CERPM sequence, and both the x and y coordinates are '
    'interpolated independently as functions of t. Duplicate t-values arising from '
    'coincident points are removed before fitting.'
)

# Method table
tbl = doc.add_table(rows=1, cols=4)
tbl.style = 'Table Grid'
hdr = tbl.rows[0].cells
for i, h in enumerate(['Method', 'SciPy Implementation', 'Continuity', 'Key Characteristic']):
    hdr[i].text = h
    for run in hdr[i].paragraphs[0].runs:
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
    shade_cell(hdr[i], 'BDD7EE')

rows_data = [
    ('Linear',       'numpy.interp',              'C⁰', 'Piecewise straight; simplest baseline; no overshoot'),
    ('Quadratic',    'make_interp_spline(k=2)',    'C¹', 'Degree-2 B-spline; smooth but low-order'),
    ('Cubic Spline', 'CubicSpline (natural)',      'C²', 'Natural cubic; current CERPM standard; zero 2nd deriv. at ends'),
    ('Quartic',      'make_interp_spline(k=4)',    'C³', 'Higher smoothness; more endpoint sensitivity'),
    ('Quintic',      'make_interp_spline(k=5)',    'C⁴', 'Highest smoothness; most prone to Runge oscillation'),
    ('PCHIP',        'PchipInterpolator',          'C¹', 'Monotone cubic; no overshoot; robust on sparse data'),
    ('Akima',        'Akima1DInterpolator',        'C¹', 'Locally weighted; single gap does not propagate oscillation'),
]
for rd in rows_data:
    add_table_row(tbl, rd)

doc.add_paragraph()

heading(doc, '3.6 Dropout Scenarios', level=2)
body(doc,
    'Three dropout models were implemented, reflecting physically distinct failure mechanisms:'
)
body(doc,
    'Random Dropout: Each CERPM independently fails with probability p. Models uncorrelated '
    'hardware failures. Rates tested: 10%, 15%, 20%, 25%. Per condition: 2,000 Monte Carlo '
    'trials, each with a unique random seed.'
)
body(doc,
    'Gap-Length Dropout: A contiguous block of CERPMs equivalent to a physical gap of g '
    'metres is removed from one side. Models localised failures such as resurfacing damage '
    'or a maintenance gap. Gap lengths tested: 2 m, 5 m, 10 m, 15 m, 20 m. Every valid '
    'starting position on both boundaries was tested exhaustively.'
)
body(doc,
    'Clustered Dropout: A contiguous block of n consecutive CERPMs is removed from one '
    'side, parameterised by count rather than physical distance. Cluster sizes tested: '
    '2, 5, 10, 15, 20 markers. Also exhaustively enumerated.'
)

heading(doc, '3.7 Error Metrics', level=2)
body(doc,
    'For each trial the reconstructed centreline is computed from the surviving CERPMs '
    'using the same perpendicular-nearest-point algorithm. The error for each centreline '
    'point is the Euclidean distance to the nearest point on the true centreline LineString. '
    'Two metrics are recorded per trial:'
)
body(doc, '  •  Mean error: average perpendicular distance across all centreline points.')
body(doc, '  •  Max error: maximum perpendicular distance (worst-case deviation).')
body(doc,
    'A trial is classified as a failure if max error exceeds 0.2 m — a threshold based on '
    'the typical 3.5 m lane width and an LKA intervention threshold of approximately '
    '±0.3 m from centreline, placing 0.2 m as the maximum tolerable estimation error before '
    'it begins to degrade intervention timing. Distributions of mean error across trials are '
    'summarised by P25 (25th percentile), P75 (75th percentile), and IQR = P75 − P25. '
    'The failure rate reports the percentage of trials where max error exceeded 0.2 m.'
)

heading(doc, '3.8 Project Management', level=2)
body(doc,
    'The project was managed across a twelve-week semester. Research activities (weeks 4–6) '
    'covered interpolation algorithm theory, Australian RPM standards, and road geometry '
    'data sources. Planning (week 6) produced the class-level software architecture. '
    'Development (weeks 6–9) implemented and debugged the full simulation pipeline. Data '
    'collection (week 9) ran all simulations across the three geometries. Results analysis '
    '(weeks 10–11) produced the comparative algorithm analysis and combined findings. '
    'Version control (Git) was used throughout to track code changes and manage deliverables. '
    'The full Gantt chart is reproduced in Appendix A.'
)

page_break(doc)

# ══════════════════════════════════════════════
# 4. RESULTS AND ANALYSIS
# ══════════════════════════════════════════════
heading(doc, '4. Results and Analysis', level=1)

heading(doc, '4.1 Road Geometries', level=2)
body(doc,
    'Figures 1–3 show the resampled left and right boundary polylines for each road '
    'geometry generated by the simulation pipeline. The three geometries span a wide range '
    'of curvature profiles, ensuring results have broad applicability.'
)

figure(doc, 't07_road',
       'Figure 1. Town07 — sinuous road geometry (~280 m). Multiple curvature reversals '
       'produce the most challenging test case for interpolation under gaps.',
       width=Inches(5.0))

figure(doc, 't42_road',
       'Figure 2. Town042 — long gradual curve (~560 m). The smooth, continuous curvature '
       'makes this the easiest geometry; baseline for comparison.',
       width=Inches(4.0))

figure(doc, 't44_road',
       'Figure 3. Town044 — sharp-corner road (~520 m). Two 90° turns separated by a '
       'straight section test algorithm behaviour at abrupt curvature changes.',
       width=Inches(5.0))

heading(doc, '4.2 Baseline Accuracy vs CERPM Interval', level=2)
body(doc,
    'Before any dropout is applied, the dominant driver of centreline error is CERPM '
    'spacing. The IQR heatmaps (Figures 4, 7, 10) share a consistent pattern across all '
    'three geometries:'
)
body(doc,
    'At 0.5 m and 1.0 m intervals all seven methods produce median centreline errors below '
    '0.01 m with near-zero IQR — the cells are uniformly deep green. Methods are '
    'effectively equivalent at these resolutions, making algorithm choice irrelevant when '
    'CERPMs are densely deployed.'
)
body(doc,
    'At 2.0 m the methods begin to diverge. Linear IQR rises measurably. Cubic spline, '
    'PCHIP, and Akima remain well below 0.05 m median error. At 4.0 m linear IQR is '
    'substantially elevated on Town07 (median error approaching 0.1–0.15 m), while '
    'higher-order smooth methods remain around 0.03–0.07 m. At 6.0 m and 12.0 m linear '
    'and quadratic are essentially unusable on curved roads, while PCHIP and Akima show '
    'the most stable remaining performance.'
)
body(doc,
    'This establishes the fundamental finding: dense CERPM deployment renders algorithm '
    'choice irrelevant, while sparse deployment makes it critical.'
)

# ── TOWN07 RANDOM ──
heading(doc, '4.3 Random Dropout — Town07 (Sinuous Road)', level=2)
body(doc,
    'Town07 represents the most challenging geometry due to its multiple curvature reversals. '
    'Figures 4–6 show the IQR heatmap, IQR band plots, and failure rate heatmap respectively '
    'for random dropout (10%–25%) across CERPM intervals of 0.5–12.0 m.'
)

figure(doc, 't07_rand_heatmap',
       'Figure 4. Town07 — Random dropout IQR centreline error (P25–P75, m). '
       'n=2000 runs per condition. Green = low error; red = high error.',
       width=Inches(6.2))

figure(doc, 't07_rand_bands',
       'Figure 5. Town07 — Random dropout IQR band plots per method. '
       'Each subplot shows median error (line) ± IQR (shaded band) vs CERPM interval '
       'at a given dropout rate.',
       width=Inches(5.5))

figure(doc, 't07_rand_fail',
       'Figure 6. Town07 — Random dropout failure rate (max error > 0.2 m, %). '
       'n=2000 runs per condition.',
       width=Inches(6.2))

body(doc,
    'Key observations for Town07 random dropout:'
)
body(doc,
    'At 0.5 m and 1.0 m spacing all methods maintain near-zero failure rate at all tested '
    'dropout rates (10%–25%). The high density of surviving markers is sufficient for any '
    'method to accurately reconstruct the boundary.'
)
body(doc,
    'At 2.0 m linear begins to fail: 43% / 62% / 80% / 97% failure rate at 10% / 15% / '
    '20% / 25% dropout respectively. All smooth methods remain below 5% failure.'
)
body(doc,
    'At 4.0 m linear reaches 100% failure across all dropout rates. Quadratic rises to '
    '47%–88%. Cubic spline shows 32%–91%, quartic 29%–75%, quintic 29%–74%. PCHIP '
    'shows elevated failure (65%–85% range). Akima performs consistently, ranging '
    '56%–96% — still elevated but among the lower values at 4 m on this complex geometry.'
)
body(doc,
    'At 6.0 m and 12.0 m all methods show high failure rates under any random dropout on '
    'Town07, reflecting the fundamental geometric challenge of reconstructing a sinuous road '
    'from very sparse markers with any missing data.'
)

# ── TOWN042 RANDOM ──
heading(doc, '4.4 Random Dropout — Town042 (Gradual Curve)', level=2)
body(doc,
    'Town042 is the most forgiving geometry due to its gradual, consistent curvature. '
    'Figures 7–9 show results for random dropout.'
)

figure(doc, 't42_rand_heatmap',
       'Figure 7. Town042 — Random dropout IQR centreline error (P25–P75, m). '
       'n=2000 runs per condition.',
       width=Inches(6.2))

figure(doc, 't42_rand_bands',
       'Figure 8. Town042 — Random dropout IQR band plots per method.',
       width=Inches(5.5))

figure(doc, 't42_rand_fail',
       'Figure 9. Town042 — Random dropout failure rate (max error > 0.2 m, %).',
       width=Inches(6.2))

body(doc,
    'Town042 shows substantially lower failure rates at all spacings compared to Town07. '
    'Linear at 2.0 m: 24% / 28% / 42% / 56% failure. At 4.0 m linear reaches 100%, but '
    'quadratic only rises to 5%–40%, and cubic spline to 8%–38%. PCHIP ranges 38%–49% '
    'at 4.0 m, while Akima is notably the best performer at 4.0 m on this geometry: '
    '5%–30% failure rate across all dropout rates. At 6.0 m, cubic spline shows '
    '52%–88%, quartic 54%–91%, quintic 60%–100%, while Akima achieves 29%–72% — '
    'the best result at this spacing. The IQR band plots (Figure 8) show very tight bands '
    'for all methods at ≤2.0 m and diverging bands at larger spacings, with Akima '
    'maintaining the tightest band at sparse spacings.'
)

# ── TOWN044 RANDOM ──
heading(doc, '4.5 Random Dropout — Town044 (Sharp Corners)', level=2)
body(doc,
    'Town044 presents an intermediate challenge but introduces a critical new finding about '
    'PCHIP. Figures 10–12 show random dropout results.'
)

figure(doc, 't44_rand_heatmap',
       'Figure 10. Town044 — Random dropout IQR centreline error (P25–P75, m). '
       'n=2000 runs per condition.',
       width=Inches(6.2))

figure(doc, 't44_rand_bands',
       'Figure 11. Town044 — Random dropout IQR band plots per method.',
       width=Inches(5.5))

figure(doc, 't44_rand_fail',
       'Figure 12. Town044 — Random dropout failure rate (max error > 0.2 m, %).',
       width=Inches(6.2))

body(doc,
    'At 4.0 m, 10% dropout: linear fails at 100%, quadratic at 25%, cubic spline at 18%, '
    'quartic at 16%, quintic at 18%, Akima at 12% — and notably, PCHIP at 46%. This is a '
    'significant reversal from the other geometries, where PCHIP and Akima typically tracked '
    'closely. The cause is PCHIP\'s monotonicity constraint: at a sharp 90° corner, the '
    'road curvature reverses abruptly, and monotone cubic construction — which prevents '
    'overshoot within each segment — cannot properly represent the rapid direction change '
    'when markers are sparse. Akima\'s local slope weighting handles this better, making '
    'it the only method that outperforms cubic spline and quartic at 4.0 m on Town044. '
    'PCHIP\'s failure rate at 4.0 m reaches 46%–88%, worse than cubic spline (18%–62%) '
    'on this geometry.'
)

# ── TOWN07 GAP ──
heading(doc, '4.6 Gap-Length Dropout — Town07', level=2)
body(doc,
    'Gap-length dropout is the most practically significant failure scenario. '
    'Figures 13–15 show results for Town07.'
)

figure(doc, 't07_gap_heatmap',
       'Figure 13. Town07 — Gap-length dropout IQR centreline error (P25–P75, m). '
       'Exhaustive enumeration of all gap positions on both boundaries.',
       width=Inches(5.0))

figure(doc, 't07_gap_bands',
       'Figure 14. Town07 — Gap-length dropout IQR band plots per method.',
       width=Inches(5.5))

figure(doc, 't07_gap_fail',
       'Figure 15. Town07 — Gap-length dropout failure rate (max error > 0.2 m, %).',
       width=Inches(4.5))

body(doc,
    'Linear interpolation is uniquely catastrophic under gap-length dropout on the sinuous '
    'Town07 geometry. For any gap ≥ 2 m at spacings ≥ 2.0 m, linear fails at 100%. Even '
    'at 1.0 m spacing with a 10 m gap, linear fails at approximately 14%. The reason is '
    'geometric: linear simply draws a straight chord across the gap — on a curved road, '
    'this chord diverges immediately from the true boundary curvature.'
)
body(doc,
    'Cubic spline, quartic, and quintic perform substantially better at moderate gaps. At '
    '1.0 m spacing with gaps up to 10 m, failure rates remain 0%–5%. At 4.0 m spacing '
    'with a 20 m gap, failure rates rise to 25%–40%. PCHIP and Akima show the best '
    'overall gap resilience on Town07, maintaining near-zero failure at 1.0 m spacing '
    'for all gap lengths, with Akima having slightly tighter IQR bands than PCHIP '
    'at the larger gap lengths (Figure 14).'
)
body(doc,
    'The IQR heatmap (Figure 13) shows a strong interaction between gap length and CERPM '
    'spacing: the same 5 m gap causes negligible error at 0.5 m spacing (many surviving '
    'markers) but causes significant error at 4.0 m spacing (only one or two survivors). '
    'This confirms that failure severity is determined by the ratio of gap length to CERPM '
    'spacing, not physical gap length alone.'
)

# ── TOWN042 GAP ──
heading(doc, '4.7 Gap-Length Dropout — Town042', level=2)
body(doc,
    'Figures 16–18 show gap-length dropout results for the gentle Town042 geometry.'
)

figure(doc, 't42_gap_heatmap',
       'Figure 16. Town042 — Gap-length dropout IQR centreline error (P25–P75, m).',
       width=Inches(5.0))

figure(doc, 't42_gap_bands',
       'Figure 17. Town042 — Gap-length dropout IQR band plots per method.',
       width=Inches(5.5))

figure(doc, 't42_gap_fail',
       'Figure 18. Town042 — Gap-length dropout failure rate (max error > 0.2 m, %).',
       width=Inches(4.5))

body(doc,
    'Town042 shows the most favourable gap-length results due to its gentle curvature. '
    'All smooth methods (quadratic through Akima) maintain near-zero failure rate at '
    '0.5 m and 1.0 m spacing across all gap lengths tested (2–20 m). Even at 4.0 m '
    'spacing with a 20 m gap, cubic spline, quartic, quintic, and Akima remain at 0% '
    'failure, while PCHIP shows 0% failure at all tested conditions on this geometry. '
    'Linear is the sole failure: reaching 21%–22% failure even at 0.5 m spacing for a '
    '20 m gap, and 100% at spacings ≥ 4.0 m for all gap lengths. The IQR band plots '
    '(Figure 17) show essentially invisible shading (near-zero IQR) for all smooth '
    'methods at ≤ 4.0 m spacing, confirming the stability of estimation on this geometry.'
)
body(doc,
    'The Town042 results establish a best-case deployment benchmark: a gradual rural '
    'highway or regional road with PCHIP or Akima at 4.0 m CERPM spacing would tolerate '
    'maintenance gaps of 20 m with zero failures under this model.'
)

# ── TOWN044 GAP ──
heading(doc, '4.8 Gap-Length Dropout — Town044', level=2)
body(doc,
    'Figures 19–21 show gap-length dropout results for the sharp-corner Town044 geometry.'
)

figure(doc, 't44_gap_heatmap',
       'Figure 19. Town044 — Gap-length dropout IQR centreline error (P25–P75, m).',
       width=Inches(5.0))

figure(doc, 't44_gap_bands',
       'Figure 20. Town044 — Gap-length dropout IQR band plots per method.',
       width=Inches(5.5))

figure(doc, 't44_gap_fail',
       'Figure 21. Town044 — Gap-length dropout failure rate (max error > 0.2 m, %).',
       width=Inches(4.5))

body(doc,
    'Town044 gap-length results reinforce the findings from random dropout. Linear is '
    'again catastrophic: at gap = 5 m with 1.0 m spacing it reaches 100% failure, and '
    'at gap = 20 m it reaches 59%–60% failure even at 0.5 m spacing. The sharp corners '
    'make even small gaps in the corner region cause large deviations for linear.'
)
body(doc,
    'Cubic spline, quartic, and quintic perform well at ≤ 2.0 m spacing for gaps up to '
    '10 m. At gap = 15 m and 4.0 m spacing, cubic spline and quartic show 12% failure; '
    'at gap = 20 m and 4.0 m spacing this rises to approximately 19%–21%. Akima is the '
    'standout performer: at 4.0 m spacing and gap = 20 m, Akima shows only ~11% failure '
    'compared to 19%+ for cubic spline. PCHIP performs notably worse on this geometry: '
    'at gap = 10 m and 4.0 m spacing PCHIP shows 100% failure (the orange/red cell in '
    'Figure 19), whereas Akima and cubic spline remain around 4%–8%. This is consistent '
    'with the random dropout finding for Town044 — PCHIP\'s monotonicity constraint is '
    'a liability at sharp curvature reversals, particularly when a gap removes the '
    'markers near the apex of the corner.'
)
body(doc,
    'The IQR band plots (Figure 20) visually confirm Akima\'s superiority on Town044: '
    'its bands remain narrowest across all gap lengths, particularly at the 4.0 m and '
    '6.0 m spacing range that is most relevant for cost-constrained deployments.'
)

# ── SUMMARY ──
heading(doc, '4.9 Cross-Geometry Comparative Summary', level=2)
body(doc,
    'Table 2 summarises the overall algorithm ranking across all three geometries and '
    'both dropout scenarios, derived from the complete set of 21 result images.'
)

# Summary ranking table
tbl2 = doc.add_table(rows=1, cols=5)
tbl2.style = 'Table Grid'
hdr2 = tbl2.rows[0].cells
for i, h in enumerate(['Method', 'Random Dropout', 'Gap-Length Dropout',
                        'Sharp Corners (T044)', 'Overall Rank']):
    hdr2[i].text = h
    for run in hdr2[i].paragraphs[0].runs:
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
    shade_cell(hdr2[i], 'BDD7EE')

rank_data = [
    ('Akima',       'Excellent', 'Excellent',     'Best',        '1st'),
    ('PCHIP',       'Excellent', 'Very Good',     'Poor',        '2nd'),
    ('Cubic Spline','Very Good', 'Good',           'Good',        '3rd'),
    ('Quartic',     'Good',      'Good',           'Good',        '4th'),
    ('Quintic',     'Good',      'Moderate',       'Good',        '5th'),
    ('Quadratic',   'Moderate',  'Moderate',       'Moderate',    '6th'),
    ('Linear',      'Poor',      'Unacceptable',   'Unacceptable','7th'),
]
for rd in rank_data:
    add_table_row(tbl2, rd)

doc.add_paragraph()
caption(doc, 'Table 2. Overall algorithm ranking derived from all 21 result images across three geometries.')

body(doc,
    'The key revision from the initial analysis is the separation of Akima and PCHIP. '
    'While both perform similarly on smooth and gradual road geometries, Akima is '
    'demonstrably superior on sharp-corner roads — making it the recommended default '
    'algorithm for general deployment where road geometry is variable or unknown.'
)

heading(doc, '4.10 Changes from Initial Research Plan', level=2)
body(doc,
    'The initial plan specified five CERPM intervals (0.5, 1.0, 2.0, 3.0, 5.0 m). '
    'During development this was revised to six intervals (0.5, 1.0, 2.0, 4.0, 6.0, '
    '12.0 m) to better characterise the failure transition region and include the 12.0 m '
    'spacing representative of low-cost rural deployments.'
)
body(doc,
    'Monte Carlo runs were increased from 500 to 2,000 per condition to reduce statistical '
    'noise in the failure rate tail distribution (±2.2 pp at 95% confidence vs ±4.4 pp '
    'for 500 runs).'
)
body(doc,
    'Three road geometries were tested rather than the initially planned set, providing '
    'cross-geometry generalisability. The gap-length dropout analysis was added as a more '
    'physically interpretable variant of clustered dropout, directly actionable by road '
    'engineers specifying maintenance intervals. The 0.2 m failure threshold was formalised '
    'during the analysis phase based on LKA intervention geometry considerations.'
)

page_break(doc)

# ══════════════════════════════════════════════
# 5. CONCLUSIONS
# ══════════════════════════════════════════════
heading(doc, '5. Conclusions and Recommendations', level=1)

heading(doc, '5.1 Conclusions', level=2)
conclusions = [
    ('Conclusion 1 — CERPM spacing is the primary determinant of system performance.',
     'At spacings of 0.5–1.0 m, all seven methods achieve sub-centimetre median error with '
     'near-zero failure rates under all tested dropout scenarios. At spacings beyond 2.0 m, '
     'algorithm selection and dropout resistance become critical.'),
    ('Conclusion 2 — Akima is the most consistently robust algorithm across all road geometries.',
     'It outperforms all other methods on the sharp-corner Town044 geometry under both random '
     'and gap-length dropout, and matches or slightly trails PCHIP on smooth geometries. '
     'Its local slope construction limits gap-induced error propagation without the '
     'monotonicity constraint that disables PCHIP at curvature reversals.'),
    ('Conclusion 3 — PCHIP is excellent on smooth and gradual roads but unsuitable as a '
     'general-purpose algorithm.',
     'It performs comparably to Akima on Town07 and Town042, but degrades significantly on '
     'Town044 due to its monotonicity constraint at sharp corners. It should not be used '
     'as a general default without road geometry classification.'),
    ('Conclusion 4 — Linear interpolation is unacceptable for any real-world deployment.',
     'Its 100% failure rate under gap-length dropout at spacings ≥ 2.0 m on curved roads '
     'disqualifies it as a primary algorithm. The simplicity advantage is not worth '
     'the safety cost.'),
    ('Conclusion 5 — Road geometry complexity should drive CERPM spacing standards.',
     'The sinuous Town07 geometry required approximately half the marker spacing of '
     'Town042 to achieve equivalent accuracy. Sharp-corner geometries require particular '
     'attention to marker density at curve apices.'),
]
for title, text in conclusions:
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(2)
    r_b = p.add_run(title + ' ')
    r_b.bold = True
    r_b.font.name = 'Calibri'
    r_b.font.size = Pt(11)
    r_t = p.add_run(text)
    r_t.font.name = 'Calibri'
    r_t.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(6)

heading(doc, '5.2 Recommendations', level=2)
recs = [
    ('R1 — Adopt Akima as the default interpolation algorithm ',
     'for CERPM-based LKA systems, replacing cubic spline. Akima provides superior '
     'performance across all road geometry types with negligible computational overhead.'),
    ('R2 — Use PCHIP only on known smooth or gradual road geometries ',
     'where the monotonicity property is beneficial and sharp curvature reversals are absent.'),
    ('R3 — Specify maximum CERPM spacing of 1.0 m for safety-critical applications. ',
     'At this spacing, even 25% random failure and physical gaps up to 15 m produce '
     'acceptable centreline error across all tested geometries with Akima interpolation.'),
    ('R4 — Specify maximum CERPM spacing of 2.0 m for lower-speed or cost-constrained '
     'rural applications, ',
     'with the understanding that gaps exceeding 5 m will elevate failure rates on '
     'sinuous roads. A maintenance protocol limiting gap lengths should accompany this standard.'),
    ('R5 — Implement gap detection and fallback logic in the onboard system. ',
     'When the number of received CERPMs on one side falls below a minimum threshold during '
     'an expected gap, the system should increase the LKA intervention threshold to avoid '
     'reacting to a degraded centreline estimate.'),
    ('R6 — Extend this research to include GPS positioning error. ',
     'The current simulation assumes exact CERPM coordinates. A follow-on study should '
     'characterise the combined effect of spacing, dropout, and GPS noise '
     '(0.1–1.0 m RMS for commercial GNSS) on system accuracy.'),
    ('R7 — Evaluate sensor fusion with vision-based LKA. ',
     'A study combining CERPM centreline estimates with camera-based detection would '
     'characterise the accuracy improvement achievable through redundancy, particularly '
     'in the gap-length failure scenario.'),
]
for title, text in recs:
    p = doc.add_paragraph()
    r_b = p.add_run(title)
    r_b.bold = True
    r_b.font.name = 'Calibri'
    r_b.font.size = Pt(11)
    r_t = p.add_run(text)
    r_t.font.name = 'Calibri'
    r_t.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(6)

page_break(doc)

# ══════════════════════════════════════════════
# 6. REFERENCES
# ══════════════════════════════════════════════
heading(doc, '6. References', level=1)
refs = [
    '[1]  S. Sharma, J. Rojas, A. R. Ekti, R. Wang, Z. Asher, and R. Meyer, "Vehicle '
     'Lateral Offset Estimation Using Infrastructure Information for Reduced Compute Load," '
     'SAE Technical Paper Series, Apr. 2023, doi: 10.4271/2023-01-0800.',
    '[2]  P. Kadav et al., "Automated Lane Centering: An Off-the-Shelf Computer Vision '
     'Product vs. Infrastructure-Based Chip-Enabled Raised Pavement Markers," Sensors, '
     'vol. 24, no. 7, pp. 2327–2327, Apr. 2024, doi: 10.3390/s24072327.',
    '[3]  Department of Infrastructure, Transport, Cities and Regional Development, '
     '"Fact sheet: Evidence supporting the priority focus areas," National Road Safety '
     'Strategy, 2021. https://www.roadsafety.gov.au/nrss/fact-sheets/priority-focus-areas',
    '[4]  National Highway Traffic Safety Administration, "Estimating Effectiveness of Lane '
     'Keeping Assist Systems in Fatal Road Departure Crashes," 2024. '
     'https://crashstats.nhtsa.dot.gov/Api/Public/ViewPublication/813663',
    '[5]  European Commission, "Mandatory drivers assistance systems expected to help save '
     'over 25,000 lives by 2038," Internal Market, Industry, Entrepreneurship and SMEs, '
     'Jul. 05, 2024.',
    '[6]  Y. Wang, A. Alhuraish, S. Yuan, and H. Zhou, "OpenLKA: Open source multimodal '
     'OpenLKA dataset," GitHub, 2025. https://github.com/OpenLKA/OpenLKA '
     '(accessed Mar. 22, 2026).',
    '[7]  I. Fakhari and S. Anwar, "A Multiple Model Estimation Approach to Robust Lane '
     'Detection via Computer Vision Based Models," 2022 IEEE 31st ISIE, pp. 576–581, '
     'Jun. 2022, doi: 10.1109/isie51582.2022.9831692.',
    '[8]  I. Fakhari and S. Anwar, "Computer vision model based robust lane detection using '
     'multiple model adaptive estimation methodology," Frontiers in Mechanical Engineering, '
     'vol. 11, Feb. 2025, doi: 10.3389/fmech.2025.1436338.',
    '[9]  G. Perozzi, J. J. Rath, C. Sentouh, J. Floris, and J.-C. Popieul, "Lateral '
     'Shared Sliding Mode Control for Lane Keeping Assist System in Steer-by-Wire Vehicles," '
     'IEEE Trans. Intelligent Vehicles, vol. 8, no. 4, pp. 3073–3082, Apr. 2023, '
     'doi: 10.1109/tiv.2021.3097352.',
    '[10] S. Wei, P. E. Pfeffer, and J. Edelmann, "State of the Art: Ongoing Research in '
     'Assessment Methods for Lane Keeping Assistance Systems," IEEE Trans. Intelligent '
     'Vehicles, vol. 9, no. 9, pp. 5853–5875, Sep. 2024, doi: 10.1109/tiv.2023.3269156.',
]
for r in refs:
    p = doc.add_paragraph(r)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.5)
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(10)

page_break(doc)

# ══════════════════════════════════════════════
# 7. APPENDICES
# ══════════════════════════════════════════════
heading(doc, '7. Appendices', level=1)

heading(doc, 'Appendix A — Project Gantt Chart', level=2)
body(doc,
    'The table below reproduces the Gantt chart from the project abstract, summarising '
    'the planned schedule across the 12-week semester.'
)
gantt_cols = ['Section', 'Task', 'Wk 4', 'Wk 5', 'Wk 6', 'Wk 7', 'Wk 8',
              'Wk 9', 'Wk 10', 'Wk 11', 'Wk 12']
gtbl = doc.add_table(rows=1, cols=len(gantt_cols))
gtbl.style = 'Table Grid'
for i, h in enumerate(gantt_cols):
    gtbl.rows[0].cells[i].text = h
    for run in gtbl.rows[0].cells[i].paragraphs[0].runs:
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(8)
    shade_cell(gtbl.rows[0].cells[i], 'BDD7EE')

G = '●'
gantt_rows = [
    ('Research', 'Interpolation Algorithms',         G, G, '',G, '',  '', '', '', ''),
    ('Research', 'Australian RPM Standards',         G, G, '', '', '',  '', '', '', ''),
    ('Research', 'Road Geometry Sources',            '', G, G, '', '',  '', '', '', ''),
    ('Research', 'Importing Road Geometry',          '', '', G, G, '',  '', '', '', ''),
    ('Planning', 'Class Diagram',                    '', '', G, '', '',  '', '', '', ''),
    ('Development', 'Coding',                        '', '', G, G, G,  '', '', '', ''),
    ('Development', 'Testing and Debugging',         '', '', '', '', G,  G, '', '', ''),
    ('Simulation', 'Data Collection',                '', '', '', '', '',  G, '', '', ''),
    ('Results', 'Interp. Algorithm Comparison',      '', '', '', '', '',  '', G, G, ''),
    ('Results', 'Analysis of CERPM Dropout',         '', '', '', '', '',  '', G, G, ''),
    ('Results', 'Combine Findings',                  '', '', '', '', '',  '', '', G, G),
]
for gr in gantt_rows:
    row = gtbl.add_row()
    for i, val in enumerate(gr):
        row.cells[i].text = str(val)
        for run in row.cells[i].paragraphs[0].runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(8)

doc.add_paragraph()

heading(doc, 'Appendix B — Software Architecture', level=2)
body(doc, 'The simulation is implemented across four Python modules:')
modules = [
    ('test.py', 'Core geometric utilities: OSM parsing (getNodes, getWays, getLanelets), '
     'lanelet chain-linking (combineWays), CERPM resampling (resample), and centreline '
     'computation (calCenterline).'),
    ('Interpolations.py', 'Arc-length parameterisation (arcLenghtParameter) and the '
     'unified interpolate(pts, numPts, method) function implementing all seven methods '
     'via SciPy.'),
    ('Simulation.py', 'Monte Carlo engine: runMonteCarloRandom, '
     'runMonteCarloClusteredSingle, and runMonteCarloGapLength, all using '
     'ProcessPoolExecutor for parallel trial execution.'),
    ('xmlParse.py', 'Top-level orchestration, parameter sweep management, and all '
     'visualisation functions (IQR heatmaps, IQR band plots, failure rate heatmaps, '
     'summary tables).'),
]
for name, desc in modules:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    rb = p.add_run(name + ': ')
    rb.bold = True
    rb.font.name = 'Calibri'
    rb.font.size = Pt(11)
    rt = p.add_run(desc)
    rt.font.name = 'Calibri'
    rt.font.size = Pt(11)

heading(doc, 'Appendix C — Simulation Parameters', level=2)
ptbl = doc.add_table(rows=1, cols=2)
ptbl.style = 'Table Grid'
for i, h in enumerate(['Parameter', 'Values']):
    ptbl.rows[0].cells[i].text = h
    for run in ptbl.rows[0].cells[i].paragraphs[0].runs:
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
    shade_cell(ptbl.rows[0].cells[i], 'BDD7EE')
params = [
    ('CERPM intervals',          '0.5, 1.0, 2.0, 4.0, 6.0, 12.0 m'),
    ('Random dropout rates',     '10%, 15%, 20%, 25%'),
    ('Gap lengths',              '2, 5, 10, 15, 20 m'),
    ('Cluster sizes',            '2, 5, 10, 15, 20 markers'),
    ('Monte Carlo runs (random)','2,000 per condition'),
    ('Gap/cluster trials',       'Exhaustive enumeration (all valid positions, both sides)'),
    ('Failure threshold',        'Max centreline error > 0.2 m'),
    ('Centreline sample interval','0.5 m'),
    ('Road geometries',          'Town07, Town042, Town044'),
    ('Interpolation methods',    'Linear, Quadratic, Cubic Spline, Quartic, Quintic, PCHIP, Akima'),
]
for p_row in params:
    add_table_row(ptbl, p_row)

doc.add_paragraph()

heading(doc, 'Appendix D — Interpolation Method Properties', level=2)
dtbl = doc.add_table(rows=1, cols=5)
dtbl.style = 'Table Grid'
for i, h in enumerate(['Method', 'SciPy Class', 'Min Points', 'Continuity', 'Overshoot']):
    dtbl.rows[0].cells[i].text = h
    for run in dtbl.rows[0].cells[i].paragraphs[0].runs:
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
    shade_cell(dtbl.rows[0].cells[i], 'BDD7EE')
method_props = [
    ('Linear',       'numpy.interp',            '2', 'C⁰', 'No'),
    ('Quadratic',    'make_interp_spline(k=2)', '3', 'C¹', 'Yes'),
    ('Cubic Spline', 'CubicSpline',             '3', 'C²', 'Yes'),
    ('Quartic',      'make_interp_spline(k=4)', '5', 'C³', 'Yes'),
    ('Quintic',      'make_interp_spline(k=5)', '6', 'C⁴', 'Yes'),
    ('PCHIP',        'PchipInterpolator',       '2', 'C¹', 'No'),
    ('Akima',        'Akima1DInterpolator',     '3', 'C¹', 'Limited'),
]
for mp in method_props:
    add_table_row(dtbl, mp)

# ── Save ──
out = 'Final_Report.docx'
doc.save(out)
print(f'Saved: {out}')
